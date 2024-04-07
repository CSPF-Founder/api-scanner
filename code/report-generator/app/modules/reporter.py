import os
import datetime

# External dependencies
from docx import Document

# Internal dependencies
from app.db.models.jobs import JobModel
from app.modules.cats import CatsReporter
from app.modules.zap import ZapReporter


class Reporter:

    def __init__(self, job: JobModel):
        self.job = job

    # Main make report function
    def run(self, *, server_url, auth_headers_data, logger) -> bool:
        scan_dir = self.job.get_local_work_dir()
        logger.info("Making Report")
        report_dir = os.path.abspath(scan_dir + "/report/")
        os.makedirs(report_dir, exist_ok=True)
        requests_dir = os.path.abspath(scan_dir + "/requests/")
        os.makedirs(requests_dir, exist_ok=True)
        document = Document()
        self.create_base(document, server_url)

        zap_scanner = ZapReporter(self.job)
        zap_scanner.add_to_report(
            document=document,
            server_url=server_url,
            auth_headers_data=auth_headers_data,
        )

        cat_scanner = CatsReporter(self.job)
        cat_scanner.add_to_report(
            document,
            logger,
        )
        document.save(report_dir + "/report.docx")
        if os.path.exists(report_dir + "/report.docx"):
            return True

        return False

    # Create base report
    def create_base(self, document, url_from_parsed):
        document.add_heading("API Security Scan of " + url_from_parsed, 0)
        report_time = datetime.datetime.now(datetime.timezone.utc).strftime(
            "%d %b %Y %H:%M UTC"
        )

        document.add_paragraph("Report generated time: " + report_time)
        document.add_page_break()
