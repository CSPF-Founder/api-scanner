import os
from urllib.parse import urlparse
import configparser
import urllib.parse

# External dependencies
from defusedxml.expatbuilder import parseString
from docx.shared import RGBColor

# Internal dependencies
from app.db.models.jobs import JobModel
from app.utils import common_utils


class ZapReporter:
    def __init__(self, job: JobModel):
        self.job = job

        self.output_dir = os.path.join(self.job.get_local_work_dir(), "zapoutput/")
        self.auth_path = os.path.join(self.output_dir, "zap_config")

    def get_auth_headers(self, auth_header_file, logger):
        """
        * Check if authfile has headers
        """
        auth_config = configparser.ConfigParser()
        auth_config.read(auth_header_file)
        list_of_values = auth_config.items("AUTH_HEADERS")
        number_of_headers = len(list_of_values)
        if number_of_headers == 1:
            return list_of_values[0]
        if number_of_headers == 0:
            logger.info("No Security Headers")
            return False
        if number_of_headers > 1:
            logger.error("Multiple auth headers. Cannot process")
            return False

    # add zap data to report
    def add_to_report(
        self,
        *,
        document,
        server_url,
        auth_headers_data,
    ):
        count = 0
        document.add_heading("Security Scan Results:", 1)
        zap_xml_path = os.path.join(self.output_dir, "zapReport.xml")
        xml_file = open(zap_xml_path, "r")
        xml_report_string = common_utils.smart_str(
            xml_file.read().replace("ZAP", "DARWIS")
        )
        if xml_report_string:
            count = self.process_xml_string(
                count=count,
                xml_report_string=xml_report_string,
                document=document,
                server_url=server_url,
                auth_headers_data=auth_headers_data,
            )

        if count == 0:
            document.add_heading("No issues found", level=2)
            document.add_page_break()

    def get_risk_from_data_node(self, data_node) -> str | None:
        try:
            risk_desc = data_node.getElementsByTagName("riskdesc")
            if not risk_desc:
                return None

            risk = risk_desc[0].firstChild.nodeValue.split()[0]
            return risk
        except Exception:
            return None

    def get_plugin_id_from_data_node(self, data_node) -> str | None:
        try:
            plugin_id = data_node.getElementsByTagName("pluginid")
            if not plugin_id:
                return None

            return plugin_id[0].firstChild.nodeValue
        except Exception:
            return None

    def get_alert_name_from_data_node(self, data_node) -> str | None:
        try:
            alert_name = data_node.getElementsByTagName("alert")
            if not alert_name:
                return None

            return alert_name[0].firstChild.nodeValue
        except Exception:
            return None

    def get_desc_from_data_node(self, data_node) -> str:
        try:
            desc = data_node.getElementsByTagName("desc")
            if not desc:
                return ""

            return desc[0].firstChild.nodeValue.replace("<p>", "").replace("</p>", "\n")

        except Exception:
            return ""

    def get_solution_from_data_node(self, data_node) -> str:
        try:
            solution = data_node.getElementsByTagName("solution")
            if not solution:
                return ""

            return (
                solution[0]
                .firstChild.nodeValue.replace("<p>", "")
                .replace("</p>", "\n")
            )
        except Exception:
            return ""

    def get_reference_from_data_node(self, data_node) -> str:
        try:
            reference = data_node.getElementsByTagName("reference")
            if not reference:
                return ""

            return (
                reference[0]
                .firstChild.nodeValue.replace("<p>", "")
                .replace("</p>", "\n")
            )
        except Exception:
            return ""

    def add_alerts_to_document(
        self,
        *,
        count: int,
        alert_node,
        document,
    ):
        high_alerts = []
        medium_alerts = []
        low_alerts = []
        for data in alert_node:
            risk_value = self.get_risk_from_data_node(data)
            if not risk_value:
                continue

            pluginid = self.get_plugin_id_from_data_node(data)

            if risk_value not in ["High", "Medium", "Low"]:
                continue

            if pluginid == "100000":
                continue

            alert_name = self.get_alert_name_from_data_node(data)
            if not alert_name:
                continue

            try:
                instances_xml = data.getElementsByTagName("instances")
            except Exception:
                instances_xml = None

            alert = {
                "alert_name": alert_name,
                "description": self.get_desc_from_data_node(data),
                "solution": self.get_solution_from_data_node(data),
                "reference": self.get_reference_from_data_node(data),
                "instances_xml": instances_xml,
                "severity": risk_value,
            }

            if risk_value == "High":
                high_alerts.append(alert)
            elif risk_value == "Medium":
                medium_alerts.append(alert)
            elif risk_value == "Low":
                low_alerts.append(alert)

        count = self.add_severity_alert_to_document(
            count=count, alerts=high_alerts, document=document
        )

        count = self.add_severity_alert_to_document(
            count=count, alerts=medium_alerts, document=document
        )

        count = self.add_severity_alert_to_document(
            count=count, alerts=low_alerts, document=document
        )

        return count

    def add_severity_alert_to_document(
        self,
        *,
        count: int,
        alerts,
        document,
    ):
        for alert in alerts:
            alert_name = alert["alert_name"]
            description = alert["description"]
            solution = alert["solution"]
            reference = alert["reference"]
            severity = alert["severity"]
            instances_xml = alert["instances_xml"]

            count = count + 1
            document.add_heading(str(count) + ". " + alert_name, level=2)
            h1 = document.add_heading("Severity" + "   :   ", level=2)
            r = h1.add_run(severity)
            if severity == "High":
                r.font.color.rgb = RGBColor(254, 0, 0)
            elif severity == "Medium":
                r.font.color.rgb = RGBColor(255, 153, 0)
            elif severity == "Low":
                r.font.color.rgb = RGBColor(0, 0, 204)
            if description:
                document.add_heading("Finding \t\t:", level=2)
                document.add_paragraph(description)
            if instances_xml:
                instances = instances_xml[0].getElementsByTagName("instance")
                i = 0
                for instance in instances:
                    if i == 0:
                        document.add_heading("URL/s \t\t:", level=2)
                    if i > 20:
                        break
                    uri = instance.getElementsByTagName("uri")
                    if not uri:
                        continue

                    uri = uri[0].firstChild
                    if not uri:
                        continue

                    uri = uri.nodeValue
                    i += 1

                    document.add_paragraph(uri)
            if solution:
                document.add_heading("Solution \t\t:", level=2)
                document.add_paragraph(solution)
            if reference:
                document.add_heading("Reference \t\t:", level=2)
                document.add_paragraph(reference)
            document.add_page_break()
        return count

    def add_remaining_alerts(
        self,
        *,
        count: int,
        alert_node,
        document,
        auth_headers_data,
    ):
        for data in alert_node:
            risk_value = self.get_risk_from_data_node(data)
            if not risk_value:
                continue

            pluginid = self.get_plugin_id_from_data_node(data)

            if risk_value != "Low" or pluginid != "100000":
                continue

            alert_name = self.get_alert_name_from_data_node(data)
            if not alert_name:
                continue

            description = self.get_desc_from_data_node(data)
            solution = self.get_solution_from_data_node(data)
            reference = self.get_reference_from_data_node(data)

            try:
                instances_xml = data.getElementsByTagName("instances")
            except Exception:
                instances_xml = None

            if pluginid == "100000" and instances_xml:
                count = count + 1
                document.add_heading(str(count) + ". " + alert_name, level=2)
                h1 = document.add_heading("Severity" + "   :   ", level=2)
                r = h1.add_run(risk_value)
                r.font.color.rgb = RGBColor(0, 0, 204)
                if description:
                    document.add_heading("Finding \t\t:", level=2)
                    document.add_paragraph(description)
                if solution:
                    document.add_heading("Solution \t\t:", level=2)
                    document.add_paragraph(solution)
                if reference:
                    document.add_heading("Reference \t\t:", level=2)
                    document.add_paragraph(reference)
                if instances_xml:
                    instances = instances_xml[0].getElementsByTagName("instance")
                    i = 0
                    for instance in instances:
                        if i == 0:
                            document.add_heading(
                                "CURL request to replicate \t\t:", level=2
                            )
                        if i > 40:
                            break
                        i += 1
                        return_curl = self.create_curl(instance, auth_headers_data)
                        document.add_paragraph(return_curl)

                document.add_page_break()
        return count

    def process_xml_string(
        self,
        *,
        count: int,
        xml_report_string,
        document,
        server_url,
        auth_headers_data,
    ):
        xml_report_string = xml_report_string.replace("\xa0", " ")
        raw_xml = parseString(common_utils.smart_str(xml_report_string))
        site_nodes = raw_xml.getElementsByTagName("site")
        for site_node in site_nodes:
            site_node_name = site_node.attributes["name"].value
            from_zap_url_hostname = urlparse(site_node_name).netloc
            from_api_url_hostname = urlparse(server_url).netloc
            if from_zap_url_hostname != from_api_url_hostname:
                continue

            alertitem_nodes = site_node.getElementsByTagName("alerts")
            for alertitem in alertitem_nodes:
                item_node = alertitem.getElementsByTagName("alertitem")
                count = self.add_alerts_to_document(
                    count=count,
                    alert_node=item_node,
                    document=document,
                )

                count = self.add_remaining_alerts(
                    count=count,
                    alert_node=item_node,
                    document=document,
                    auth_headers_data=auth_headers_data,
                )
        return count

    # Create CURL for zap request
    def create_curl(self, xml_node, auth_header):
        if auth_header:
            curl_auth_header = " -H '" + auth_header[0] + ": " + auth_header[1] + "'"
        else:
            curl_auth_header = ""
        try:
            uri = xml_node.getElementsByTagName("uri")[0].firstChild.nodeValue
        except Exception:
            uri = ""
        try:
            method = xml_node.getElementsByTagName("method")[0].firstChild.nodeValue
        except Exception:
            method = ""
        try:
            param = xml_node.getElementsByTagName("param")[0].firstChild.nodeValue
        except Exception:
            param = ""
        try:
            attack = xml_node.getElementsByTagName("attack")[0].firstChild.nodeValue
        except Exception:
            attack = ""
        if param != "" and attack != "":
            param_str = urllib.parse.quote(param)
            attack_str = urllib.parse.quote(attack)
            curl_data = f" --data '{param_str}={attack_str}'"
        else:
            curl_data = ""

        curl_string = (
            f"curl -X '{method}' --url "
            f"'{uri}'"
            " -H 'accept: application/json' " + curl_data + curl_auth_header
        )

        return curl_string
